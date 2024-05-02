import os
import PyPDF2
import docx
from googletrans import Translator
from flask import Flask, request, jsonify

app = Flask(__name__)

@app.route('/', methods=['GET' , 'POST'])
def translate_document():
    # Get user's selected target language
    target_language = request.form.get('target_language')

    # Check if a file is uploaded
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    # Check if the file is empty
    if file.filename == '':
        return jsonify({'error': 'Empty file uploaded'}), 400

    # Check if the file is supported
    filename, file_extension = os.path.splitext(file.filename)
    if file_extension.lower() not in ['.pdf', '.doc', '.docx']:
        return jsonify({'error': 'Unsupported file format'}), 400

    try:
        # Extract text from PDF
        if file_extension.lower() == '.pdf':
            pdf_reader = PyPDF2.PdfFileReader(file)
            text = ''
            for page_num in range(pdf_reader.numPages):
                text += pdf_reader.getPage(page_num).extractText()

        # Extract text from DOCX
        elif file_extension.lower() == '.docx':
            docx_doc = docx.Document(file)
            text = '\n'.join([paragraph.text for paragraph in docx_doc.paragraphs])

        # Extract text from DOC
        elif file_extension.lower() == '.doc':
            # Convert DOC to DOCX using LibreOffice (requires LibreOffice to be installed)
            os.system(f'libreoffice --headless --convert-to docx {file.filename}')
            docx_doc = docx.Document(f'{filename}.docx')
            text = '\n'.join([paragraph.text for paragraph in docx_doc.paragraphs])
            os.remove(f'{filename}.docx')  # Remove the temporary DOCX file

        # Translate the extracted text
        translator = Translator()
        translated_text = translator.translate(text, dest=target_language).text

        return jsonify({'translated_text': translated_text}), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
