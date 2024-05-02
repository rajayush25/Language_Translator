# from googletrans import Translator

# def translate_text(text, target_language):
#     translator = Translator()
#     translated = translator.translate(text, dest=target_language)
#     return translated.text

# # Example usage
# input_text = "Hello, how are you?"
# desired_language = 'te'  # 'es' for Spanish, you can change it to any desired language code

# translated_text = translate_text(input_text, desired_language)
# print(f"Translated text: {translated_text}")
from flask import Flask, request, jsonify , render_template
from texttranlstionfinal import translate_text
import os
import PyPDF2
import docx
from googletrans import Translator
import cv2
import pytesseract
from PIL import Image
import tkinter as tk


app = Flask(__name__)

def translate_text(text, target_language):
    translator = Translator()
    translated = translator.translate(text, dest=target_language)
    return translated.text

@app.route('/', methods=['GET' , 'POST'])
def handle_translation():
    if request.method == 'POST':
        text = request.form.get("text")
        destlang = request.form.get("destlang")
        text_trans = translate_text(text , destlang)
        return render_template("translated_text.html" , text_trans = text_trans)
    return render_template("index.html")

@app.route("/doc_trans" , methods=["POST"])
def doc_trans():
    target_language = request.form.get('target_language')
    # Check if a file is uploaded
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    # Check if the file is empty
    if file.filename == '':
        return jsonify({'error': 'Empty file uploaded'}), 400

    # Check if the file is supported
    filename, file_extension = os.path.splitext(file.filename)
    if file_extension.lower() not in ['.pdf', '.doc', '.docx']:
        return jsonify({'error': 'Unsupported file format'}), 400

    try:
        # Extract text from PDF
        if file_extension.lower() == '.pdf':
            pdf_reader = PyPDF2.PdfFileReader(file)
            text = ''
            for page_num in range(pdf_reader.numPages):
                text += pdf_reader.getPage(page_num).extractText()

        # Extract text from DOCX
        elif file_extension.lower() == '.docx':
            docx_doc = docx.Document(file)
            text = '\n'.join([paragraph.text for paragraph in docx_doc.paragraphs])

        # Extract text from DOC
        elif file_extension.lower() == '.doc':
            # Convert DOC to DOCX using LibreOffice (requires LibreOffice to be installed)
            os.system(f'libreoffice --headless --convert-to docx {file.filename}')
            docx_doc = docx.Document(f'{filename}.docx')
            text = '\n'.join([paragraph.text for paragraph in docx_doc.paragraphs])
            os.remove(f'{filename}.docx')  # Remove the temporary DOCX file

        # Translate the extracted text
        translator = Translator()
        translated_text = translator.translate(text, dest=target_language).text
        return render_template("trans_doc.html" , translated_text = translated_text  )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route("/img_trans" , methods=['GET' , "POST"])
def img_trans():
    def ocr_image(image_path):
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text

# Function to translate text from one language to another using Google Translate API
    def translate_text(text, src_lang, dest_lang):
        translator = Translator()
        translated = translator.translate(text, src=src_lang, dest=dest_lang)
        return translated.text

# Function to capture an image using the system's camera
    def capture_image():
        camera = cv2.VideoCapture(0)
        return_value, image = camera.read()
        cv2.imwrite('captured_image.jpg', image)
        del camera
        return 'captured_image.jpg'

# Function triggered when translation button is clicked
    def translate():
        global translated_text
        # Capture an image using the system's camera
        image_path = capture_image()
        
        # Perform OCR on the captured image to extract text
        extracted_text = ocr_image(image_path)

        # Get the selected language from the dropdown
        selected_lang = language_var.get()

        # Translate the extracted text to the selected language
        if selected_lang == 'Marathi':
            translated_text = translate_text(extracted_text, 'en', 'mr')
        elif selected_lang == 'Gujarati':
            translated_text = translate_text(extracted_text, 'en', 'gu')
        elif selected_lang == 'Hindi':
            translated_text = translate_text(extracted_text, 'en', 'hi')
        elif selected_lang == 'Bengali':
            translated_text = translate_text(extracted_text, 'en', 'bn')
        elif selected_lang == 'Punjabi':
            translated_text = translate_text(extracted_text, 'en', 'pa')
        elif selected_lang == 'Tamil':
            translated_text = translate_text(extracted_text, 'en', 'ta')
        elif selected_lang == 'Telugu':
            translated_text = translate_text(extracted_text, 'en', 'te')
        elif selected_lang == 'Kannada':
            translated_text = translate_text(extracted_text, 'en', 'kn')
        elif selected_lang == 'Malayalam':
            translated_text = translate_text(extracted_text, 'en', 'ml')
        elif selected_lang == 'Odia':
            translated_text = translate_text(extracted_text, 'en', 'or')
        elif selected_lang == 'Urdu':
            translated_text = translate_text(extracted_text, 'en', 'ur')
        elif selected_lang == 'Sanskrit':
            translated_text = translate_text(extracted_text, 'en', 'sa')

        # Display translated text in the text box
        text_box.delete(1.0, tk.END)
        text_box.insert(tk.END, translated_text)

    # GUI setup
    root = tk.Tk()
    root.title("Image Translator")

    # Language selection dropdown
    language_var = tk.StringVar(root)
    language_var.set("Select Language")

    language_dropdown = tk.OptionMenu(root, language_var, "Marathi", "Gujarati", "Hindi", "Bengali", "Punjabi", "Tamil", "Telugu", "Kannada", "Malayalam", "Odia", "Urdu", "Sanskrit")
    language_dropdown.pack()

    # Translated text display
    text_box = tk.Text(root, height=10, width=50)
    text_box.pack()

    # Translate button
    translate_button = tk.Button(root, text="Translate", command=translate)
    translate_button.pack()
    root.mainloop()
    return render_template("index.html")
if __name__ == '__main__':
    app.run(debug=True)
